# CREAR REPORTE DE ANALISIS DE DATOS FINAL

from docx import Document
from docx.shared import Inches
import pandas as pd
import os

print("📄 Generando reporte Word...")

# ===============================
# Cargar datos
# ===============================
df = pd.read_csv("ventas_limpias.xlsx")

# ===============================
# Crear documento
# ===============================
doc = Document()

# Título
doc.add_heading("Reporte de Análisis de Ventas", level=1)

doc.add_paragraph(
    "Este reporte presenta un análisis exploratorio de las ventas, "
    "incluyendo limpieza de datos, análisis temporal, distribución por "
    "tienda y productos con mayor impacto en ingresos."
)

# ===============================
# Resumen general
# ===============================
doc.add_heading("1. Resumen General", level=2)

doc.add_paragraph(f"Total de registros analizados: {len(df):,}")
doc.add_paragraph(f"Ventas totales: ${df['Item_Outlet_Sales'].sum():,.2f}")
doc.add_paragraph(f"Venta promedio por registro: ${df['Item_Outlet_Sales'].mean():,.2f}")

# ===============================
# Tabla Top 10 Productos
# ===============================
doc.add_heading("2. Top 10 Productos por Ingresos", level=2)

top_10 = (
    df.groupby("Item_Identifier")["Item_Outlet_Sales"]
    .sum()
    .sort_values(ascending=False)
    .head(10)
    .reset_index()
)

table = doc.add_table(rows=1, cols=2)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Producto"
hdr_cells[1].text = "Ingresos Totales"

for _, row in top_10.iterrows():
    row_cells = table.add_row().cells
    row_cells[0].text = row["Item_Identifier"]
    row_cells[1].text = f"${row['Item_Outlet_Sales']:,.2f}"

# ===============================
# Gráficos
# ===============================
doc.add_heading("3. Visualizaciones", level=2)

graficos = [
    "ventas_por_mes.png",
    "distribucion_ventas.png",
    "ventas_por_tienda.png",
    "ventas_por_tamano.png",
    "top_10_productos.png"
]

for grafico in graficos:
    if os.path.exists(grafico):
        doc.add_paragraph(grafico.replace("_", " ").replace(".png", "").title())
        doc.add_picture(grafico, width=Inches(5))

# ===============================
# Conclusiones
# ===============================
doc.add_heading("4. Conclusiones", level=2)

doc.add_paragraph(
    "El análisis permitió identificar patrones de ventas relevantes, "
    "productos con mayor impacto económico y diferencias significativas "
    "entre tipos de tiendas. La visualización temporal evidencia variaciones "
    "claras en el comportamiento de ventas."
)

# ===============================
# Guardar documento
# ===============================
nombre_archivo = "Reporte_Analisis_Ventas.docx"
doc.save(nombre_archivo)

print(f"✅ Reporte generado exitosamente: {nombre_archivo}")

print("📂 Buscando archivos en:", os.path.abspath(output_ventas))
print("📄 Existe CSV:", os.path.exists(os.path.join(output_ventas, "ventas_limpias.xlsx")))


